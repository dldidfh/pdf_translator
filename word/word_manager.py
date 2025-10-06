import numpy as np
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --------- 설정값 (필요시 미세조정) ----------
FONT_NAME = "D2Coding"          # 모노스페이스 + 한글 지원 권장 (미설치면 Courier New 등으로 교체)
FONT_SIZE_PT = 11               # 워드 글자 크기
DPI = 96                        # OCR bbox가 픽셀 기준일 때 화면/이미지 DPI
CHAR_WIDTH_PT_RATIO = 0.60      # '한 글자 폭 ≈ 0.60 * font_size_pt' 가정값 (폰트별로 0.55~0.62 범위 튜닝)
LINE_JOIN_TOL = 0.6             # 줄 묶기 허용 오차(글자 높이*계수)
TREAT_CJK_AS_DOUBLE = False     # CJK(한글/한자)를 2칸으로 볼지 여부 (모노스페이스 폰트면 False 권장)
# --------------------------------------------

def to_int64(v):
    """xywh가 int16이어도 안전하게 int64로 승격"""
    if isinstance(v, np.ndarray):
        return v.astype(np.int64)
    elif isinstance(v, (np.integer, int)):
        return np.int64(v)
    elif isinstance(v, (list, tuple)):
        return np.array(v, dtype=np.int64)
    else:
        raise TypeError(f"unsupported type: {type(v)}")

def px_per_char(font_pt=FONT_SIZE_PT, ratio=CHAR_WIDTH_PT_RATIO, dpi=DPI):
    """한 글자 폭(픽셀) 추정치: (pt*ratio)/72 inch * dpi"""
    return (font_pt * ratio / 72.0) * dpi

def char_count(text: str):
    """문자 개수(모노스페이스 가정). CJK를 2폭으로 볼지 옵션."""
    if not TREAT_CJK_AS_DOUBLE:
        return len(text)
    cnt = 0
    for ch in text:
        # 간단 판정: ASCII이면 1, 그 외(대부분 CJK)는 2
        cnt += 1 if ord(ch) < 128 else 2
    return cnt

def group_lines_xywh(items):
    """
    items: [{'text': str, 'xywh': [x,y,w,h]}]  (픽셀, 원점: 좌상단)
    세로 위치로 줄을 묶고, 줄 내에서는 x순으로 정렬
    """
    # int64 승격
    norm = []
    for text, xyxy in items:
        x1, y1, x2, y2 = xyxy
        w = x2 - x1 
        h = y2 - y1
        xywh = [x1,y1,w,h]
        xywh = to_int64(xywh)
        x, y, w, h = xywh.tolist()
        norm.append({"text": text, "x": int(x), "y": int(y), "w": int(w), "h": int(h)})

    # 높이의 중앙값으로 라인 임계값 산정
    hs = [r["h"] for r in norm if r["h"] > 0]
    med_h = np.median(hs) if hs else 20
    line_tol = max(4, int(med_h * LINE_JOIN_TOL))

    # y-top 기준으로 정렬 후 클러스터링
    norm.sort(key=lambda t: t["y"])
    lines = []
    cur = []
    last_y = None
    for t in norm:
        if last_y is None or abs(t["y"] - last_y) <= line_tol:
            cur.append(t)
            last_y = t["y"] if last_y is None else min(last_y, t["y"])
        else:
            lines.append(cur)
            cur = [t]
            last_y = t["y"]
    if cur:
        lines.append(cur)

    # 각 라인 내 x순 정렬
    for ln in lines:
        ln.sort(key=lambda t: t["x"])
    return lines

def apply_mono_font(run, name=FONT_NAME, size_pt=FONT_SIZE_PT):
    run.font.name = name
    run.font.size = Pt(size_pt)
    # 동아시아 폰트(한글)도 동일 폰트로 강제
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)


def layout_with_spaces(doc: Document, ocr_items, font_scale=0.2):
    """
    ocr_items: [{'text': str, 'xywh': [x, y, w, h]}]
    OCR bbox 정보를 기반으로 자동 줄 정렬 + 줄바꿈 + 높이에 따른 폰트 크기 조절.
    """
    ppc = px_per_char()
    lines = group_lines_xywh(ocr_items)

    prev_y = None
    prev_h = None

    for line in lines:
        ys = [t["y"] for t in line]
        hs = [t["h"] for t in line if t["h"] > 0]
        cur_y = min(ys)
        cur_h = np.median(hs) if hs else 0

        # (1) 줄간 간격에 따른 빈 줄 추가
        if prev_y is not None and prev_h is not None:
            gap = cur_y - (prev_y + prev_h)
            if gap > prev_h * 1.1:
                n_blank = int(round(gap / max(prev_h, 1)))
                for _ in range(n_blank - 1):
                    doc.add_paragraph("")

        # (2) 줄 텍스트 구성
        cur_cols = 0
        buf = []
        for tok in line:
            target_cols = int(round(tok["x"] / ppc))
            spaces_needed = max(0, target_cols - cur_cols)
            buf.append(" " * spaces_needed)
            buf.append(tok["text"])
            cur_cols = target_cols + char_count(tok["text"])

        vis = "".join(buf).replace(" ", "\u00A0")

        # (3) 폰트 크기 = OCR bbox 높이에 비례
        dynamic_font_size_pt = max(6, cur_h * 72 / DPI * font_scale)

        p = doc.add_paragraph()
        run = p.add_run(vis)
        apply_mono_font(run, name=FONT_NAME, size_pt=dynamic_font_size_pt)

        prev_y = cur_y
        prev_h = cur_h
